import requests, json, sys, tempfile, os, uuid, datetime, time

ts = datetime.datetime.now().strftime("%H%M%S%f")
suf = uuid.uuid4().hex[:12]
ADMIN_EMAIL = "admin@example.com"
REC_EMAIL   = "re2e_%s_%s@example.com" % (ts, suf)
CAN_EMAIL   = "ce2e_%s_%s@example.com" % (ts, suf)
BASE = "https://cv-match-platform-api-tk2.azurewebsites.net"
ok = lambda name, cond: results.append(f"  PASS {name}") if cond else results.append(f"  FAIL {name}")
results = []

def req(method, path, body=None, token=None):
    h = {}
    if token: h["Authorization"] = "Bearer " + token
    try:
        r = requests.request(method, BASE + path, json=body, headers=h, timeout=20)
        if r.status_code >= 400: return None
        return r.json() if r.text else {}
    except: return None

# 1
ok("Health", (req("GET","/health") or {}).get("status")=="ok")
# 2-4 Register/Login. Public register allows candidate/recruiter only; admin uses seeded account.
adminT = (req("POST","/api/auth/login", {"email":ADMIN_EMAIL,"password":"Admin123!"}) or {}).get("access_token")
ok("Login seeded admin", bool(adminT))
rec   = req("POST","/api/auth/register", {"email":REC_EMAIL,"password":"Rec123!","name":"E2E R","role":"recruiter"})
ok("Register recruiter", rec and rec.get("user",{}).get("id"))
can   = req("POST","/api/auth/register", {"email":CAN_EMAIL,"password":"Can123!","name":"E2E C","role":"candidate"})
ok("Register candidate", can and can.get("user",{}).get("id"))
recT   = (req("POST","/api/auth/login", {"email":REC_EMAIL,"password":"Rec123!"}) or {}).get("access_token")
canT   = (req("POST","/api/auth/login", {"email":CAN_EMAIL,"password":"Can123!"}) or {}).get("access_token")
ok("Login recruiter", bool(recT))
ok("Login candidate", bool(canT))
if not (adminT and recT and canT):
    for r in results: print(r)
    print("FATAL: login failed"); sys.exit(1)

# 5
ok("Admin /me role=admin", (req("GET","/api/auth/me", token=adminT) or {}).get("role")=="admin")
ok("Recruiter /me role=recruiter", (req("GET","/api/auth/me", token=recT) or {}).get("role")=="recruiter")
ok("Candidate /me role=candidate", (req("GET","/api/auth/me", token=canT) or {}).get("role")=="candidate")

# 6 Create job
job = req("POST","/api/jobs", {"title":"Backend Engineer","company_name":"E2E","location":"HCM","required_skills":["python","fastapi","mongodb","docker","azure"],"required_experience":4,"description":"APIs","status":"open"}, token=recT)
jid = (job or {}).get("id")
ok("Create job", bool(jid))
ok("Job detail", (req("GET","/api/jobs/%s" % jid) or {}).get("title")=="Backend Engineer") if jid else ok("Job detail", False)
ok("List jobs", req("GET","/api/jobs") is not None)
ok("My jobs", req("GET","/api/jobs/my", token=recT) is not None)

# 7 Upload CV
from docx import Document
tmp = os.path.join(tempfile.gettempdir(), "cv_%s.docx" % ts)
doc = Document()
doc.add_paragraph("E2E Can")
doc.add_paragraph("Email: %s" % CAN_EMAIL)
doc.add_paragraph("Phone: 0912345678")
doc.add_paragraph("Skills: Python, FastAPI, MongoDB, Docker, Azure")
doc.add_paragraph("Experience: 4 years")
doc.save(tmp)
cid = None
try:
    with open(tmp,"rb") as f:
        r = requests.post(BASE+"/api/cvs/upload", headers={"Authorization":"Bearer "+canT}, files={"file":("cv.docx",f,"application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, timeout=30)
        cid = (r.json().get("id") if r.status_code < 400 and r.text and isinstance(r.json(), dict) else None)
        ok("Upload CV", r.status_code < 400)
except Exception as e: ok("Upload CV", False)
cv_ready = False
for _ in range(45):
    time.sleep(2)
    uploaded = req("GET","/api/cvs/%s" % cid, token=canT) if cid else None
    if uploaded and uploaded.get("processing_status") == "done":
        cv_ready = True
        break

cvs = req("GET","/api/cvs/my", token=canT)
ok("My CVs", cvs is not None)
cv_list = cvs if isinstance(cvs, list) else (cvs.get("cvs",[]) if isinstance(cvs,dict) else [])
cvD = None
if cid:
    cvD = req("GET","/api/cvs/%s" % cid, token=canT)
    ok("CV detail", bool(cvD))
    data = (cvD or {}).get("extracted_data", {})
    ok("CV has email", bool(data.get("email") or cvD.get("email")))

# 8 Apply
app = None
if jid:
    app = req("POST","/api/applications", {"job_id":jid, "cv_id":cid}, token=canT)
    app_id = (app or {}).get("id")
    ok("Apply job", bool(app_id))
else: app_id = None
ok("My applications", req("GET","/api/applications/my", token=canT) is not None)
if jid: ok("Recruiter view applications", req("GET","/api/applications/job/%s" % jid, token=recT) is not None)
if app_id: ok("Update status", bool((req("PUT","/api/applications/%s/status" % app_id, {"status":"shortlisted"}, token=recT) or {}).get("status")=="shortlisted"))

# 9 Matching
ok("Run matching", req("POST","/api/matches/run", {"job_id":jid}, token=recT) is not None)
if jid: ok("Matches for job", req("GET","/api/matches/job/%s" % jid, token=recT) is not None)
ok("My matches", req("GET","/api/matches/my", token=canT) is not None)

# 10 Conversations + Messages
conv = req("POST","/api/conversations", {"participant_id": (can or {}).get("user",{}).get("id"), "job_id": jid}, token=recT)
conv_id = (conv or {}).get("id")
ok("Create conversation", bool(conv_id))
ok("My conversations", req("GET","/api/conversations/my", token=recT) is not None)
if conv_id: ok("Send message", bool(req("POST","/api/messages", {"conversation_id":conv_id,"content":"Hi"}, token=recT)))
if conv_id: ok("Read messages recruiter", req("GET","/api/messages/%s" % conv_id, token=recT) is not None)
if conv_id: ok("Read messages candidate", req("GET","/api/messages/%s" % conv_id, token=canT) is not None)

# 11 Profile update
ok("Profile update", bool(req("PUT","/api/users/me", {"name":"Updated"}, token=canT)))

# 12 Admin
ok("Admin list users", req("GET","/api/admin/users", token=adminT) is not None)
ok("Admin stats", req("GET","/api/admin/stats", token=adminT) is not None)
ok("Admin list CVs", req("GET","/api/admin/cvs", token=adminT) is not None)
ok("Admin list jobs", req("GET","/api/admin/jobs", token=adminT) is not None)

t = len(results); p = sum(1 for r in results if "PASS" in r); f = t - p
print()
print("=== PRODUCTION A-Z E2E ===")
print("Backend: " + BASE)
for r in results: print(r)
print()
print("=== %d PASS / %d FAIL (%d total) ===" % (p, f, t))
if f: sys.exit(1)

