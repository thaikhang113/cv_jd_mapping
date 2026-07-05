import asyncio
from pathlib import Path
from bson import ObjectId
from docx import Document
from app.services import cv_worker


class Cursor:
    def __init__(self, rows):
        self.rows = list(rows)

    def sort(self, key, direction):
        reverse = direction < 0
        self.rows.sort(key=lambda row: row.get(key, 0), reverse=reverse)
        return self

    def __aiter__(self):
        self.index = 0
        return self

    async def __anext__(self):
        if self.index >= len(self.rows):
            raise StopAsyncIteration
        row = self.rows[self.index]
        self.index += 1
        return row


class Collection:
    def __init__(self, rows=None):
        self.rows = rows or []

    async def find_one(self, query):
        return next((row for row in self.rows if all(row.get(k) == v for k, v in query.items())), None)

    def find(self, query=None):
        query = query or {}
        return Cursor([row for row in self.rows if all(row.get(k) == v for k, v in query.items())])

    async def update_one(self, query, update, upsert=False):
        row = await self.find_one(query)
        if not row and upsert:
            row = {"_id": ObjectId(), **query}
            self.rows.append(row)
        if row:
            row.update(update.get("$set", {}))
            for key in update.get("$unset", {}):
                row.pop(key, None)


class FakeDB:
    def __init__(self, cv, job):
        self.cvs = Collection([cv])
        self.jobs = Collection([job])
        self.matching_results = Collection([])


def test_process_queue_item_extracts_and_matches(monkeypatch, tmp_path):
    cv_id, owner_id, job_id, recruiter_id = ObjectId(), ObjectId(), ObjectId(), ObjectId()
    path = tmp_path / "cv.docx"
    doc = Document()
    doc.add_paragraph("Candidate One")
    doc.add_paragraph("Python FastAPI MongoDB Docker 3 years Remote")
    doc.save(path)
    cv = {"_id": cv_id, "owner_id": owner_id, "file_path": str(path), "raw_text": "", "extracted_data": {}}
    job = {"_id": job_id, "recruiter_id": recruiter_id, "status": "open", "title": "Backend", "description": "FastAPI", "required_skills": ["python", "fastapi"], "required_experience": 2, "location": "Remote"}
    fake_db = FakeDB(cv, job)
    monkeypatch.setattr(cv_worker, "db", fake_db)

    matched_jobs = asyncio.run(cv_worker.process_queue_item({"cv_id": cv_id}))

    assert matched_jobs == 1
    assert cv["processing_status"] == "done"
    assert "python" in cv["extracted_data"]["skills"]
    assert fake_db.matching_results.rows[0]["overall_score"] > 60


def test_process_queue_item_marks_cv_done_before_best_effort_matching(monkeypatch, tmp_path):
    cv_id, owner_id = ObjectId(), ObjectId()
    path = tmp_path / "cv.docx"
    doc = Document()
    doc.add_paragraph("Candidate Two")
    doc.add_paragraph("Email: two@example.com")
    doc.add_paragraph("Python FastAPI 2 years Remote")
    doc.save(path)
    cv = {"_id": cv_id, "owner_id": owner_id, "uploaded_by_role": "candidate", "file_path": str(path), "raw_text": "", "extracted_data": {}}
    bad_job = {"_id": ObjectId(), "status": "open", "title": "Broken", "description": "Python", "required_skills": ["python"]}
    fake_db = FakeDB(cv, bad_job)
    fake_db.users = Collection([{"_id": owner_id}])
    monkeypatch.setattr(cv_worker, "db", fake_db)

    matched_jobs = asyncio.run(cv_worker.process_queue_item({"cv_id": cv_id}))

    assert matched_jobs == 0
    assert cv["processing_status"] == "done"
    assert cv["extracted_data"]["email"] == "two@example.com"
    assert cv["matching_error"]
    assert fake_db.users.rows[0]["primary_cv_id"] == cv_id

def test_enqueue_cv_can_preserve_done_status(monkeypatch):
    cv_id, owner_id = ObjectId(), ObjectId()
    cv = {"_id": cv_id, "owner_id": owner_id, "processing_status": "done"}
    fake_db = FakeDB(cv, {})
    fake_db.cv_processing_queue = Collection([])
    monkeypatch.setattr(cv_worker, "db", fake_db)

    asyncio.run(cv_worker.enqueue_cv(cv_id, owner_id, mark_queued=False))

    assert cv["processing_status"] == "done"
    assert fake_db.cv_processing_queue.rows[0]["status"] == "pending"
