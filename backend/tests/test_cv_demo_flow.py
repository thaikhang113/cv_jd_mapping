import asyncio
from bson import ObjectId
from pymongo.errors import DuplicateKeyError
from app.routes import applications, matches
from app.services import cv_worker
from app.services.cv_parser import parse_cv_text
from app.schemas.common import ApplicationIn


class Result:
    def __init__(self, inserted_id): self.inserted_id = inserted_id


class Cursor:
    def __init__(self, rows): self.rows = list(rows)
    def sort(self, key, direction): self.rows.sort(key=lambda r: r.get(key, 0), reverse=direction < 0); return self
    def __aiter__(self): self.i = 0; return self
    async def __anext__(self):
        if self.i >= len(self.rows): raise StopAsyncIteration
        row = self.rows[self.i]; self.i += 1; return row


class Collection:
    def __init__(self, rows=None, unique=None): self.rows = rows or []; self.unique = unique
    def _match(self, row, query):
        for k, v in query.items():
            if isinstance(v, dict) and "$in" in v:
                if row.get(k) not in v["$in"]: return False
            elif row.get(k) != v: return False
        return True
    async def find_one(self, query, *args, sort=None, **kwargs):
        rows = [r for r in self.rows if self._match(r, query)]
        if sort and rows:
            key, direction = sort[0]; rows.sort(key=lambda r: r.get(key, 0), reverse=direction < 0)
        return rows[0] if rows else None
    def find(self, query=None, *args, **kwargs): return Cursor([r for r in self.rows if self._match(r, query or {})])
    async def insert_one(self, doc):
        if self.unique and any(all(row.get(k) == doc.get(k) for k in self.unique) for row in self.rows): raise DuplicateKeyError("dup")
        doc = {"_id": ObjectId(), **doc}; self.rows.append(doc); return Result(doc["_id"])
    async def update_one(self, query, update, upsert=False):
        row = await self.find_one(query)
        if not row and upsert:
            row = {"_id": ObjectId(), **query}; self.rows.append(row)
        if row:
            row.update(update.get("$set", {}))
            for k, v in update.get("$inc", {}).items(): row[k] = row.get(k, 0) + v


def test_parse_cv_reads_contact_and_vietnamese_skill_heading():
    text = """
    Nguyen Van A 1783255350398
    Email: example@gmail.com
    Phone: 0912 345 678
    Ky nang
    React, Node.js, MongoDB, Python, FastAPI
    Cong nghe
    Docker, Azure, SQL
    """
    data = parse_cv_text(text)
    assert data["email"] == "example@gmail.com"
    assert data["phone"] == "0912 345 678"
    assert {"react", "node", "mongodb", "python", "fastapi", "docker", "azure", "sql"}.issubset(set(data["skills"]))


def test_processed_candidate_cv_becomes_primary(monkeypatch, tmp_path):
    from docx import Document
    user_id, cv_id = ObjectId(), ObjectId()
    path = tmp_path / "cv.docx"
    doc = Document(); doc.add_paragraph("Candidate One"); doc.add_paragraph("candidate@example.com 0912345678 Python React MongoDB 3 years"); doc.save(path)
    cv = {"_id": cv_id, "owner_id": user_id, "uploaded_by_role": "candidate", "file_path": str(path), "raw_text": "", "extracted_data": {}}
    user = {"_id": user_id, "role": "candidate"}
    fake_db = type("DB", (), {})()
    fake_db.cvs = Collection([cv]); fake_db.users = Collection([user]); fake_db.jobs = Collection([]); fake_db.matching_results = Collection([])
    monkeypatch.setattr(cv_worker, "db", fake_db)
    asyncio.run(cv_worker.process_queue_item({"cv_id": cv_id}))
    assert user["primary_cv_id"] == cv_id
    assert cv["extracted_data"]["email"] == "candidate@example.com"
    assert cv["extracted_data"]["phone"] == "0912345678"


def test_reapply_updates_existing_application_to_current_cv(monkeypatch):
    user_id, recruiter_id, job_id, old_cv_id, new_cv_id, app_id = ObjectId(), ObjectId(), ObjectId(), ObjectId(), ObjectId(), ObjectId()
    fake_db = type("DB", (), {})()
    fake_db.jobs = Collection([{"_id": job_id, "status": "open", "recruiter_id": recruiter_id}])
    fake_db.cvs = Collection([{"_id": new_cv_id, "owner_id": user_id, "processing_status": "done"}, {"_id": old_cv_id, "owner_id": user_id, "processing_status": "done"}])
    fake_db.applications = Collection([{"_id": app_id, "job_id": job_id, "candidate_id": user_id, "cv_id": old_cv_id, "recruiter_id": recruiter_id, "status": "applied"}], unique=["job_id", "candidate_id"])
    fake_db.matching_results = Collection([])
    monkeypatch.setattr(applications, "db", fake_db)
    row = asyncio.run(applications.apply(ApplicationIn(job_id=str(job_id)), {"id": str(user_id), "primary_cv_id": str(new_cv_id)}))
    assert row["id"] == str(app_id)
    assert row["cv_id"] == str(new_cv_id)


def test_application_enrichment_includes_cv_contact(monkeypatch):
    candidate_id, recruiter_id, job_id, cv_id = ObjectId(), ObjectId(), ObjectId(), ObjectId()
    row = {"_id": ObjectId(), "job_id": job_id, "candidate_id": candidate_id, "recruiter_id": recruiter_id, "cv_id": cv_id, "status": "applied"}
    fake_db = type("DB", (), {})()
    fake_db.jobs = Collection([{"_id": job_id, "title": "Backend", "company_name": "Demo"}])
    fake_db.users = Collection([{"_id": candidate_id, "name": "Candidate"}, {"_id": recruiter_id, "name": "Recruiter"}])
    fake_db.cvs = Collection([{"_id": cv_id, "filename": "cv.pdf", "extracted_data": {"name": "CV Name", "email": "cv@example.com", "phone": "0912345678", "skills": ["python"]}}])
    monkeypatch.setattr(applications, "db", fake_db)
    out = asyncio.run(applications.enrich_applications([row]))[0]
    assert out["cv_email"] == "cv@example.com"
    assert out["cv_phone"] == "0912345678"
    assert out["cv_skills"] == ["python"]


def test_run_matches_includes_cv_contact_and_sorted(monkeypatch):
    recruiter_id, job_id = ObjectId(), ObjectId()
    cv1, cv2 = ObjectId(), ObjectId()
    job = {"_id": job_id, "recruiter_id": recruiter_id, "status": "open", "title": "Backend", "description": "Python FastAPI", "required_skills": ["python", "fastapi"], "required_experience": 1, "location": "Remote"}
    fake_db = type("DB", (), {})()
    fake_db.jobs = Collection([job])
    fake_db.cvs = Collection([
        {"_id": cv1, "owner_id": ObjectId(), "processing_status": "done", "raw_text": "Python FastAPI Remote 3 years", "extracted_data": {"name": "Good", "email": "good@example.com", "phone": "1", "skills": ["python", "fastapi"], "experience_years": 3, "location": "Remote"}},
        {"_id": cv2, "owner_id": ObjectId(), "processing_status": "done", "raw_text": "Excel", "extracted_data": {"name": "Weak", "email": "weak@example.com", "phone": "2", "skills": ["excel"], "experience_years": 0}},
    ])
    fake_db.matching_results = Collection([])
    monkeypatch.setattr(matches, "db", fake_db)
    rows = asyncio.run(matches.run_matches({"job_id": str(job_id)}, {"id": str(recruiter_id), "role": "recruiter"}))
    assert rows[0]["overall_score"] >= rows[1]["overall_score"]
    assert rows[0]["cv_email"] == "good@example.com"
    assert rows[0]["cv_phone"] == "1"
    assert "python" in rows[0]["cv_skills"]

def test_extract_phone_does_not_include_location_suffix():
    data = parse_cv_text("Candidate\nEmail: a@example.com | Phone: 0912345678 | Ho Chi Minh City")
    assert data["phone"] == "0912345678"
